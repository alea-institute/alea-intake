"""DOCX export adapter using python-docx with legal formatting.

Converts Markdown to a DOCX document with Times New Roman 12pt,
headers/footers, TOC placeholder, and Markdown heading hierarchy.
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.services.output.export.base import ExportAdapter
from app.services.output.schemas import OutputContext, OutputProfile


class DOCXAdapter(ExportAdapter):
    """Export adapter producing DOCX via python-docx with legal formatting."""

    async def export(
        self,
        markdown: str,
        context: OutputContext,
        profile: OutputProfile,
    ) -> bytes:
        """Convert Markdown to DOCX with legal formatting.

        1. Create Document with Times New Roman 12pt default
        2. Add header with matter_title, footer with confidentiality notice
        3. Add TOC placeholder
        4. Parse Markdown and convert to DOCX elements
        5. Save to bytes

        Args:
            markdown: Rendered Markdown content.
            context: Output data for metadata.
            profile: Output profile for branding.

        Returns:
            DOCX bytes.
        """
        doc = Document()

        # Set default font
        self._set_default_font(doc, profile)

        # Set header and footer
        self._set_header_footer(doc, context.matter_title)

        # Add TOC placeholder
        self._add_toc_placeholder(doc)

        # Parse and add Markdown content
        if markdown:
            self._parse_markdown_to_docx(doc, markdown)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @property
    def content_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    @property
    def file_extension(self) -> str:
        return "docx"

    @staticmethod
    def _set_default_font(doc: Document, profile: OutputProfile) -> None:
        """Set document default font to Times New Roman 12pt (or org branding font)."""
        font_name = "Times New Roman"
        if profile.org_branding and profile.org_branding.font_name:
            font_name = profile.org_branding.font_name

        style = doc.styles["Normal"]
        font = style.font
        font.name = font_name
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)

    @staticmethod
    def _set_header_footer(doc: Document, matter_title: str) -> None:
        """Set document header to matter title and footer to confidentiality notice."""
        section = doc.sections[0]

        # Header
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = matter_title
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = "Confidential - Attorney Work Product"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def _add_toc_placeholder(doc: Document) -> None:
        """Insert a Table of Contents field code placeholder.

        Uses the XML workaround to insert a TOC field that Word will update
        when the document is opened.
        """
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(instr_text)

        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_char_separate)

        # Placeholder text shown before TOC is updated
        toc_text = OxmlElement("w:t")
        toc_text.text = "[Table of Contents — Update in Word]"
        run._r.append(toc_text)

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)

        # Page break after TOC
        doc.add_page_break()

    def _parse_markdown_to_docx(self, doc: Document, markdown: str) -> None:
        """Parse Markdown content line by line and convert to DOCX elements.

        Supports: headings (#), bullet lists (-), numbered lists (1.),
        bold (**), italic (*), horizontal rules (---), and plain paragraphs.
        """
        lines = markdown.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            # Horizontal rule -> page break
            if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
                doc.add_page_break()
                i += 1
                continue

            # Headings
            heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                doc.add_heading(text, level=level)
                i += 1
                continue

            # Bullet list
            bullet_match = re.match(r"^[-*]\s+(.*)", stripped)
            if bullet_match:
                text = bullet_match.group(1)
                para = doc.add_paragraph(style="List Bullet")
                self._add_formatted_text(para, text)
                i += 1
                continue

            # Numbered list
            num_match = re.match(r"^\d+\.\s+(.*)", stripped)
            if num_match:
                text = num_match.group(1)
                para = doc.add_paragraph(style="List Number")
                self._add_formatted_text(para, text)
                i += 1
                continue

            # Plain paragraph
            para = doc.add_paragraph()
            self._add_formatted_text(para, stripped)
            i += 1

    @staticmethod
    def _add_formatted_text(paragraph, text: str) -> None:
        """Add text to a paragraph with inline bold and italic formatting.

        Handles **bold** and *italic* markers.
        """
        # Split by bold/italic markers
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                if part:
                    paragraph.add_run(part)
