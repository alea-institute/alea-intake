"""DOCX text extraction using python-docx.

Extracts text with structural elements (headings, paragraphs, tables).
Runs synchronous python-docx operations in a thread executor to avoid blocking.
"""

from __future__ import annotations

import asyncio
from pathlib import Path


def _extract_sync(file_path: Path) -> tuple[str, list[dict]]:
    """Synchronous DOCX extraction using python-docx.

    Classifies elements based on paragraph style names:
    - Styles starting with "Heading" -> element_type="heading"
    - Everything else -> element_type="paragraph"
    - Table content -> element_type="table"

    Returns:
        Tuple of (full_text, elements) where elements are dicts with
        text, element_type, and optional heading_level keys.
    """
    from docx import Document

    doc = Document(str(file_path))
    elements: list[dict] = []
    text_parts: list[str] = []

    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            # Extract heading level from style name (e.g., "Heading 1" -> 1)
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            elements.append({
                "text": text,
                "element_type": "heading",
                "heading_level": level,
            })
        else:
            elements.append({
                "text": text,
                "element_type": "paragraph",
            })

        text_parts.append(text)

    # Process tables
    for table in doc.tables:
        rows_text: list[str] = []
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells]
            rows_text.append("\t".join(cell_texts))

        table_text = "\n".join(rows_text)
        if table_text.strip():
            elements.append({
                "text": table_text,
                "element_type": "table",
            })
            text_parts.append(table_text)

    full_text = "\n\n".join(text_parts)
    return full_text, elements


async def extract_docx(file_path: Path) -> tuple[str, list[dict]]:
    """Extract text and structural elements from a DOCX file.

    Runs synchronous python-docx operations in a thread executor to avoid
    blocking the event loop.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Tuple of (full_text, elements) where elements are dicts with
        text, element_type, and optional heading_level keys.
    """
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _extract_sync, file_path)
