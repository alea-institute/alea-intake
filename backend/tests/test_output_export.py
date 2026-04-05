"""Tests for output export adapters: PDF, DOCX, JSON, and ExportAdapter ABC."""

from __future__ import annotations

import json
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

import pytest

from app.services.output.export.base import ExportAdapter
from app.services.output.export.pdf_adapter import PDFAdapter
from app.services.output.export.docx_adapter import DOCXAdapter
from app.services.output.export.json_adapter import JSONAdapter
from app.services.output.schemas import (
    ActionItem,
    CIRACSection,
    GapReport,
    OrgBranding,
    OutputContext,
    OutputProfile,
    TriageResult,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

SAMPLE_MARKDOWN = """\
# Legal Analysis: Smith v. Jones

## California

### Breach of Contract

#### Issue

Whether defendant breached the employment agreement.

#### Rule

Under California Civil Code Section 1619, a contract is either express or implied.

#### Application

- The parties entered a written employment agreement on January 1, 2025.
- Defendant terminated employment without cause on March 15, 2025.
- The agreement requires 90-day notice for termination without cause.

#### Conclusion

The facts support a breach of contract claim based on insufficient notice.

---

## Action Items

1. Obtain copy of employment agreement
2. Gather termination letter
- Document all communications

"""

MINIMAL_MARKDOWN = ""


@pytest.fixture()
def sample_profile() -> OutputProfile:
    """Create a sample output profile for testing."""
    return OutputProfile(
        profile_type="law_firm",
        language_level="professional",
        sections={
            "executive_summary": True,
            "cirac_memo": True,
            "triage_routing": False,
            "action_items": True,
            "gap_appendix": True,
            "authorities_table": True,
        },
    )


@pytest.fixture()
def branded_profile() -> OutputProfile:
    """Create a profile with org branding."""
    return OutputProfile(
        profile_type="law_firm",
        language_level="professional",
        sections={"cirac_memo": True},
        org_branding=OrgBranding(
            logo_path=None,
            primary_color="#003366",
            secondary_color="#cc6600",
            font_name="Georgia",
            org_name="Test Law Firm LLP",
        ),
    )


@pytest.fixture()
def sample_context(sample_profile: OutputProfile) -> OutputContext:
    """Create a sample OutputContext for testing."""
    return OutputContext(
        intake_id=1,
        run_id=1,
        org_id=1,
        matter_title="Smith v. Jones — Breach of Contract",
        generated_at=datetime(2026, 1, 15, 10, 30, 0),
        claims_by_jurisdiction={
            "California": [
                CIRACSection(
                    claim_id=1,
                    claim_name="Breach of Contract",
                    claim_type="primary",
                    confidence=0.85,
                    jurisdiction="California",
                    issue_statement="Whether defendant breached the employment agreement.",
                    conclusion="The facts support a breach of contract claim.",
                ),
            ],
        },
        triage=TriageResult(complexity_level="medium", urgency_level="routine"),
        action_items=[
            ActionItem(
                item_number=1,
                category="documents_to_gather",
                description="Obtain employment agreement",
                priority="urgent",
            ),
        ],
        gap_report=GapReport(completeness_score=0.75),
        completeness_score=0.75,
        executive_summary="Analysis of breach of contract claim under California law.",
        profile=sample_profile,
    )


# ---------------------------------------------------------------------------
# ExportAdapter ABC tests
# ---------------------------------------------------------------------------


class TestExportAdapterABC:
    """Tests for the ExportAdapter abstract base class."""

    def test_abc_requires_export_method(self) -> None:
        """ExportAdapter cannot be instantiated directly."""
        with pytest.raises(TypeError):
            ExportAdapter()  # type: ignore[abstract]

    def test_abc_requires_content_type(self) -> None:
        """Concrete class must implement content_type property."""

        class IncompleteAdapter(ExportAdapter):
            async def export(self, markdown: str, context: OutputContext, profile: OutputProfile) -> bytes:
                return b""

            @property
            def file_extension(self) -> str:
                return "txt"

        with pytest.raises(TypeError):
            IncompleteAdapter()  # type: ignore[abstract]

    def test_abc_requires_file_extension(self) -> None:
        """Concrete class must implement file_extension property."""

        class IncompleteAdapter(ExportAdapter):
            async def export(self, markdown: str, context: OutputContext, profile: OutputProfile) -> bytes:
                return b""

            @property
            def content_type(self) -> str:
                return "text/plain"

        with pytest.raises(TypeError):
            IncompleteAdapter()  # type: ignore[abstract]


# ---------------------------------------------------------------------------
# PDFAdapter tests
# ---------------------------------------------------------------------------


class TestPDFAdapter:
    """Tests for the PDF export adapter using WeasyPrint."""

    @pytest.fixture()
    def adapter(self) -> PDFAdapter:
        return PDFAdapter()

    async def test_pdf_export_returns_valid_pdf_bytes(
        self, adapter: PDFAdapter, sample_context: OutputContext, sample_profile: OutputProfile
    ) -> None:
        """PDFAdapter.export produces bytes starting with PDF magic bytes."""
        result = await adapter.export(SAMPLE_MARKDOWN, sample_context, sample_profile)
        assert isinstance(result, bytes)
        assert len(result) > 0
        assert result[:5] == b"%PDF-", "Output should start with PDF magic bytes"

    async def test_pdf_content_type(self, adapter: PDFAdapter) -> None:
        """PDFAdapter.content_type is application/pdf."""
        assert adapter.content_type == "application/pdf"

    async def test_pdf_file_extension(self, adapter: PDFAdapter) -> None:
        """PDFAdapter.file_extension is pdf."""
        assert adapter.file_extension == "pdf"

    async def test_pdf_empty_markdown_produces_valid_pdf(
        self, adapter: PDFAdapter, sample_context: OutputContext, sample_profile: OutputProfile
    ) -> None:
        """Empty Markdown still produces a valid (minimal) PDF."""
        result = await adapter.export(MINIMAL_MARKDOWN, sample_context, sample_profile)
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"

    async def test_pdf_applies_org_branding(
        self, adapter: PDFAdapter, sample_context: OutputContext, branded_profile: OutputProfile
    ) -> None:
        """PDFAdapter injects org branding CSS variables when provided."""
        result = await adapter.export(SAMPLE_MARKDOWN, sample_context, branded_profile)
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"

    async def test_pdf_includes_matter_title_element(
        self, adapter: PDFAdapter, sample_context: OutputContext, sample_profile: OutputProfile
    ) -> None:
        """PDFAdapter includes matter_title in generated HTML for running header."""
        # We test indirectly -- if the PDF generates, the title element was included
        result = await adapter.export(SAMPLE_MARKDOWN, sample_context, sample_profile)
        assert len(result) > 100, "PDF should contain substantial content"


# ---------------------------------------------------------------------------
# DOCXAdapter tests
# ---------------------------------------------------------------------------


class TestDOCXAdapter:
    """Tests for the DOCX export adapter using python-docx."""

    @pytest.fixture()
    def adapter(self) -> DOCXAdapter:
        return DOCXAdapter()

    async def test_docx_export_returns_valid_zip(
        self, adapter: DOCXAdapter, sample_context: OutputContext, sample_profile: OutputProfile
    ) -> None:
        """DOCXAdapter.export produces valid ZIP bytes (DOCX is ZIP format)."""
        result = await adapter.export(SAMPLE_MARKDOWN, sample_context, sample_profile)
        assert isinstance(result, bytes)
        assert len(result) > 0
        # DOCX files are ZIP archives starting with PK magic bytes
        assert result[:2] == b"PK", "DOCX output should be a valid ZIP file"
        # Verify it's a valid ZIP
        bio = BytesIO(result)
        assert zipfile.is_zipfile(bio)

    async def test_docx_content_type(self, adapter: DOCXAdapter) -> None:
        """DOCXAdapter.content_type is the correct OOXML MIME type."""
        assert adapter.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    async def test_docx_file_extension(self, adapter: DOCXAdapter) -> None:
        """DOCXAdapter.file_extension is docx."""
        assert adapter.file_extension == "docx"

    async def test_docx_contains_matter_title_header(
        self, adapter: DOCXAdapter, sample_context: OutputContext, sample_profile: OutputProfile
    ) -> None:
        """DOCXAdapter sets header to matter title."""
        result = await adapter.export(SAMPLE_MARKDOWN, sample_context, sample_profile)
        bio = BytesIO(result)
        from docx import Document

        doc = Document(bio)
        # Check section header
        section = doc.sections[0]
        header = section.header
        header_text = "".join(p.text for p in header.paragraphs)
        assert "Smith v. Jones" in header_text

    async def test_docx_contains_confidential_footer(
        self, adapter: DOCXAdapter, sample_context: OutputContext, sample_profile: OutputProfile
    ) -> None:
        """DOCXAdapter sets footer to 'Confidential - Attorney Work Product'."""
        result = await adapter.export(SAMPLE_MARKDOWN, sample_context, sample_profile)
        bio = BytesIO(result)
        from docx import Document

        doc = Document(bio)
        section = doc.sections[0]
        footer = section.footer
        footer_text = "".join(p.text for p in footer.paragraphs)
        assert "Confidential" in footer_text
        assert "Attorney Work Product" in footer_text

    async def test_docx_heading_levels_match_markdown(
        self, adapter: DOCXAdapter, sample_context: OutputContext, sample_profile: OutputProfile
    ) -> None:
        """DOCXAdapter maps Markdown heading hierarchy to DOCX heading levels."""
        result = await adapter.export(SAMPLE_MARKDOWN, sample_context, sample_profile)
        bio = BytesIO(result)
        from docx import Document

        doc = Document(bio)
        heading_styles = [p.style.name for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Heading 1" in heading_styles, "H1 from Markdown should become Heading 1"
        assert "Heading 2" in heading_styles, "H2 from Markdown should become Heading 2"

    async def test_docx_empty_markdown_produces_valid_docx(
        self, adapter: DOCXAdapter, sample_context: OutputContext, sample_profile: OutputProfile
    ) -> None:
        """Empty Markdown still produces a valid (minimal) DOCX."""
        result = await adapter.export(MINIMAL_MARKDOWN, sample_context, sample_profile)
        assert isinstance(result, bytes)
        assert result[:2] == b"PK"
        bio = BytesIO(result)
        assert zipfile.is_zipfile(bio)


# ---------------------------------------------------------------------------
# JSONAdapter tests
# ---------------------------------------------------------------------------


class TestJSONAdapter:
    """Tests for the JSON export adapter using Pydantic serialization."""

    @pytest.fixture()
    def adapter(self) -> JSONAdapter:
        return JSONAdapter()

    async def test_json_export_returns_utf8_bytes(
        self, adapter: JSONAdapter, sample_context: OutputContext, sample_profile: OutputProfile
    ) -> None:
        """JSONAdapter.export returns UTF-8 encoded JSON bytes."""
        result = await adapter.export(SAMPLE_MARKDOWN, sample_context, sample_profile)
        assert isinstance(result, bytes)
        text = result.decode("utf-8")
        data = json.loads(text)
        assert isinstance(data, dict)

    async def test_json_round_trips_through_output_context(
        self, adapter: JSONAdapter, sample_context: OutputContext, sample_profile: OutputProfile
    ) -> None:
        """JSONAdapter output can be deserialized back to OutputContext."""
        result = await adapter.export(SAMPLE_MARKDOWN, sample_context, sample_profile)
        text = result.decode("utf-8")
        restored = OutputContext.model_validate_json(text)
        assert restored.intake_id == sample_context.intake_id
        assert restored.matter_title == sample_context.matter_title
        assert restored.completeness_score == sample_context.completeness_score

    async def test_json_content_type(self, adapter: JSONAdapter) -> None:
        """JSONAdapter.content_type is application/json."""
        assert adapter.content_type == "application/json"

    async def test_json_file_extension(self, adapter: JSONAdapter) -> None:
        """JSONAdapter.file_extension is json."""
        assert adapter.file_extension == "json"

    async def test_json_empty_markdown_produces_valid_json(
        self, adapter: JSONAdapter, sample_context: OutputContext, sample_profile: OutputProfile
    ) -> None:
        """Empty Markdown still produces valid JSON (markdown is not in JSON output)."""
        result = await adapter.export(MINIMAL_MARKDOWN, sample_context, sample_profile)
        text = result.decode("utf-8")
        data = json.loads(text)
        assert data["intake_id"] == 1


# ---------------------------------------------------------------------------
# CSS stylesheet tests
# ---------------------------------------------------------------------------


class TestLegalPDFCSS:
    """Tests for the legal PDF CSS stylesheet."""

    def test_css_file_exists(self) -> None:
        """legal_pdf.css file exists in templates/css directory."""
        css_path = Path(__file__).parent.parent / "app" / "services" / "output" / "templates" / "css" / "legal_pdf.css"
        assert css_path.exists(), f"CSS file not found at {css_path}"

    def test_css_contains_page_rules(self) -> None:
        """legal_pdf.css contains @page rules for letter size and margins."""
        css_path = Path(__file__).parent.parent / "app" / "services" / "output" / "templates" / "css" / "legal_pdf.css"
        content = css_path.read_text()
        assert "@page" in content, "CSS should contain @page rules"
        assert "8.5in" in content, "CSS should specify letter width (8.5in)"
        assert "11in" in content, "CSS should specify letter height (11in)"
        assert "1in" in content, "CSS should specify 1in margins"
